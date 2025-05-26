def convert_docx_to_xlsx_improved(docx_file_path, xlsx_file_path):
    from docx import Document
    import pandas as pd
    import re

    doc = Document(docx_file_path)
    
    # Collect all key-value pairs from all tables
    data_dict = {}
    
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ').replace('\t', ' ') for cell in row.cells]
            
            # Skip completely empty rows
            if not any(cell.strip() for cell in cells):
                continue
            
            # Skip header rows (like "TERMS OF ISSUE")
            if len(cells) >= 2 and cells[1] and cells[1].upper() == "TERMS OF ISSUE":
                continue
            
            # Process rows with at least 2 columns
            if len(cells) >= 2:
                key = ""
                value = ""
                
                # Case 1: Key in second column, value in third column
                if cells[1].strip() and len(cells) > 2:
                    key = cells[1].strip()
                    value = cells[2].strip() if cells[2].strip() else "Not specified"
                
                # Case 2: Key in first column, value in second column
                elif cells[0].strip() and cells[1].strip():
                    key = cells[0].strip()
                    value = cells[1].strip()
                
                # Case 3: Only second column has content (use as both key and value)
                elif cells[1].strip() and not cells[0].strip() and (len(cells) <= 2 or not cells[2].strip()):
                    key = cells[1].strip()
                    value = "Yes" if key else ""
                
                # Clean up the key and add to dictionary if valid
                if key and key not in ["", "TERMS OF ISSUE"]:
                    # Clean up key name - remove extra spaces and special characters
                    clean_key = re.sub(r'\s+', ' ', key).strip()
                    
                    # Handle duplicate keys by adding suffix
                    original_key = clean_key
                    counter = 1
                    while clean_key in data_dict:
                        clean_key = f"{original_key}_{counter}"
                        counter += 1
                    
                    data_dict[clean_key] = value
    
    #splitting date to seperate day, month and year
    from datetime import datetime

    # Try to extract and parse the Issue Opening Date
    issue_date_raw = data_dict.get("Issue Opening Date")
    if issue_date_raw:
        try:
            # Convert formats like 'April 9, 2025' to datetime object
            parsed_date = datetime.strptime(issue_date_raw, "%B %d, %Y")            
            data_dict['date_ddmmyyyy'] = parsed_date.strftime("%d%m%Y")
            #converting to individual d d m m y y y y cells
            data_dict['d1'] = data_dict['date_ddmmyyyy'][0]
            data_dict['d2'] = data_dict['date_ddmmyyyy'][1]
            data_dict['m1'] = data_dict['date_ddmmyyyy'][2]
            data_dict['m2'] = data_dict['date_ddmmyyyy'][3]
            data_dict['y1'] = data_dict['date_ddmmyyyy'][4]
            data_dict['y2'] = data_dict['date_ddmmyyyy'][5]
            data_dict['y3'] = data_dict['date_ddmmyyyy'][6]
            data_dict['y4'] = data_dict['date_ddmmyyyy'][7]
        
        
        except ValueError as e:
            print(f"⚠️ Could not parse Issue Opening Date: {issue_date_raw}")

    # Create DataFrame
    if data_dict:
        df = pd.DataFrame([data_dict])
    else:
        df = pd.DataFrame()
    # Write to Excel with formatting
    with pd.ExcelWriter(xlsx_file_path, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name="Terms", index=False)
        
        # Get workbook and worksheet objects for formatting
        workbook = writer.book
        worksheet = writer.sheets['Terms']
        
        # Add formatting
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'bg_color': '#D7E4BC'
        })
        
        cell_format = workbook.add_format({
            'text_wrap': True,
            'valign': 'top'
        })
        
        # Apply header formatting
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
        
        # Apply cell formatting and adjust column widths
        for col_num, column in enumerate(df.columns):
            max_length = max(
                df[column].astype(str).map(len).max(),  # Max length in column
                len(str(column))  # Length of column name
            )
            # Set column width (with some padding)
            worksheet.set_column(col_num, col_num, min(max_length + 2, 50))
        
        # Set row height for better readability
        worksheet.set_row(0, 30)  # Header row
        for row_num in range(1, len(df) + 1):
            worksheet.set_row(row_num, 20)
    writer.close()
    print(f"✅ Export complete!")
    print(f"📊 Total fields extracted: {len(data_dict)}")
    print(f"📄 Saved to: {xlsx_file_path}")
    
    return data_dict

# Alternative function for better debugging and validation
# Enhanced function for better table structure handling
def convert_docx_to_xlsx(docx_file_path, xlsx_file_path, processing_date=None, debug=True):
    from docx import Document
    import pandas as pd
    import re
    from datetime import datetime

    doc = Document(docx_file_path)
    
    data_dict = {}
    debug_info = []
    
    for table_idx, table in enumerate(doc.tables):
        if debug:
            print(f"\n📋 Processing Table {table_idx + 1}:")
        
        for row_idx, row in enumerate(table.rows):
            # Clean and extract all cell contents
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ').replace('\t', ' ')
                # Remove excessive whitespace
                cell_text = re.sub(r'\s+', ' ', cell_text).strip()
                cells.append(cell_text)
            
            # Skip completely empty rows
            if not any(cell.strip() for cell in cells):
                continue
            
            # Skip header rows (containing "TERMS OF ISSUE")
            if any("TERMS OF ISSUE" in cell.upper() for cell in cells):
                continue
            
            # Initialize extraction variables
            key = ""
            value = ""
            extraction_method = ""
            
            # Handle different column structures
            num_cols = len(cells)
            non_empty_cells = [cell for cell in cells if cell.strip()]
            
            if debug:
                print(f"  Row {row_idx + 1}: {num_cols} columns, {len(non_empty_cells)} non-empty")
                for i, cell in enumerate(cells):
                    if cell.strip():
                        print(f"    Col {i+1}: '{cell[:50]}{'...' if len(cell) > 50 else ''}'")
            
            # Strategy 1: Find key-value pairs in adjacent columns
            if num_cols >= 2:
                for i in range(num_cols - 1):
                    if cells[i].strip() and cells[i+1].strip():
                        # Only skip exact matches of these phrases
                        skip_phrases = [
                            'terms of issue',
                            'terms and conditions',
                            '>>',
                            '***',
                            '*****'
                        ]
                        # Don't skip if it's a valid header
                        valid_headers = [
                            'issue price',
                            'issue opening date',
                            'issue closing date',
                            'discount at which security is issued'
                        ]
                        cell_lower = cells[i].lower()
                        
                        if (not any(phrase in cell_lower for phrase in skip_phrases) or 
                            any(header in cell_lower for header in valid_headers)):
                            key = cells[i].strip()
                            value = cells[i+1].strip()
                            extraction_method = f"Col{i+1}→Col{i+2}"
                            break
            
           # Strategy 2: Handle single column with meaningful content
            if not key and num_cols >= 1:
                skip_phrases = [
                    'terms of issue',
                    'terms and conditions',
                    '***',
                    '>>>'
                ]
                valid_headers = [
                    'issue price',
                    'issue opening date',
                    'issue closing date',
                    'discount at which security is issued'
                ]
                
                for i, cell in enumerate(cells):
                    cell_lower = cell.strip().lower()
                    if (cell.strip() and 
                        not any(skip in cell_lower for skip in skip_phrases) or
                        any(header in cell_lower for header in valid_headers)):
                        if len(cell.strip()) > 3:  # Avoid very short strings
                            key = cell.strip()  # Use actual content instead of Field_X_Y_Z
                            value = cell.strip()
                            extraction_method = f"Single_Col{i+1}"
                            break
            
            # Strategy 3: Handle multi-column data (3+ columns)
            if not key and num_cols >= 3:
                if debug:
                    print("\n→ Entering Strategy 3")
                    print(f"    Number of columns: {num_cols}")
                    print(f"    Cells content: {cells}")

                # Ensure we have cells to work with
                if cells and len(cells) > 0:
                    # Check if this is a Coupon row
                    first_cell = cells[0].strip().lower() if cells[0] else ""
                    print(f"    First cell content: '{first_cell}'")
                    
                    if first_cell == 'coupon':
                        print("    ✓ Found Coupon row!")
                        
                        # First pass: Store Coupon → Condition (Col1 → Col2)
                        coupon_key = 'Coupon' if 'Coupon' not in data_dict else 'Coupon_1'
                        condition_text = cells[1].strip() if len(cells) > 1 else ""
                        data_dict[coupon_key] = condition_text
                        
                        if debug:
                            print(f"    ✓ First pass - {coupon_key}: {condition_text[:50]}...")
                        
                        # Second pass: Store Condition → Formula (Col2 → Col3)
                        formula_key = condition_text  # Use the condition as the key
                        formula_value = cells[2].strip() if len(cells) > 2 else ""
                        
                        # Store formula with condition as key
                        if formula_key and formula_value:
                            data_dict[formula_key] = formula_value
                            if debug:
                                print(f"    ✓ Second pass - Formula stored: {formula_value}")
                        
                        # Set method for debug info
                        extraction_method = "Coupon_Special"
                        
                        if debug:
                            print("    ✓ Coupon row processing complete")
                        
                        key = None  # Prevent further processing of this row
                        continue
                    else:
                        if debug:
                            print("    → Not a Coupon row, continuing normal processing")
                                        
            # Store the data if we found a valid key-value pair
            if key and key.strip():
                clean_key = key.strip()
                clean_value = value.strip() if value else "Present"
                
                # Handle duplicate keys
                original_key = clean_key
                counter = 1
                while clean_key in data_dict:
                    clean_key = f"{original_key}_{counter}"
                    counter += 1
                
                data_dict[clean_key] = clean_value
                if processing_date:
                    date_components = handle_processing_date(processing_date)
                    if date_components:
                        data_dict.update(date_components)
                        if debug:
                            print(f"📅 Added processing date: {processing_date}")

                        
                if debug:
                    debug_info.append({
                        'Table': table_idx + 1,
                        'Row': row_idx + 1,
                        'Method': extraction_method,
                        'Key': clean_key,
                        'Value': clean_value[:100] + '...' if len(clean_value) > 100 else clean_value,
                        'Original_Cells': ' | '.join([f"Col{i+1}: {cell}" for i, cell in enumerate(cells) if cell.strip()])
                    })
                    print(f"    ✓ Extracted: {clean_key} = {clean_value[:50]}{'...' if len(clean_value) > 50 else ''}")

    # Process special date fields
    date_fields_to_process = [
        "Issue Opening Date", "Issue Closing Date", "Pay-in-Date", 
        "Date of Allotment", "Deemed Date of Allotment", "Initial Fixing Date",
        "Final Fixing Date", "Redemption Date", "Coupon / Dividend payment dates"
    ]
    
    for date_field in date_fields_to_process:
        date_value = data_dict.get(date_field)
        if date_value:
            try:
                # Try different date formats
                date_formats = ["%B %d, %Y", "%B %d,%Y", "%d %B %Y", "%d/%m/%Y", "%Y-%m-%d"]
                parsed_date = None
                
                for fmt in date_formats:
                    try:
                        parsed_date = datetime.strptime(date_value, fmt)
                        break
                    except ValueError:
                        continue
                
                if parsed_date:
                    date_key = f"{date_field.replace(' ', '_').lower()}_formatted"
                    data_dict[date_key] = parsed_date.strftime("%d%m%Y")
                    
                    # For Issue Opening Date, create individual digit fields
                    if date_field == "Issue Opening Date":
                        formatted_date = parsed_date.strftime("%d%m%Y")
                        data_dict['d1'] = formatted_date[0]
                        data_dict['d2'] = formatted_date[1]
                        data_dict['m1'] = formatted_date[2]
                        data_dict['m2'] = formatted_date[3]
                        data_dict['y1'] = formatted_date[4]
                        data_dict['y2'] = formatted_date[5]
                        data_dict['y3'] = formatted_date[6]
                        data_dict['y4'] = formatted_date[7]
                        
                        if debug:
                            print(f"📅 Processed {date_field}: {date_value} → {formatted_date}")
                
            except Exception as e:
                if debug:
                    print(f"⚠️ Could not parse date field {date_field}: {date_value} - {e}")
    # Extract series number from Product Code
    if 'Product Code' in data_dict:
        series_number = extract_series_number(data_dict['Product Code'])
        data_dict['series_number'] = series_number
        if debug:
            print(f"📎 Extracted series number: {series_number}")
    # Extract issue size number from Issue Size
    if 'Issue Size' in data_dict:
        issue_size_num, formatted_amount = extract_issue_size_number(data_dict['Issue Size'])
        data_dict['issue_size_number'] = issue_size_num
        data_dict['total_amount'] = formatted_amount
        if debug:
            print(f"💰 Extracted issue size number: {issue_size_num}")
            print(f"💰 Extracted total amount: {formatted_amount}")
    # Extract tenor days number
    if 'Tenor In Days' in data_dict:
        tenor_days = extract_tenor_days_number(data_dict['Tenor In Days'])
        data_dict['tenor_days_num'] = tenor_days
        if debug:
            print(f"📅 Extracted tenor days: {tenor_days}")
    # Extract face value number
    if 'Face Value' in data_dict:
        face_value_num, formatted_face_value = extract_face_value_number(data_dict['Face Value'])
        data_dict['face_value_num'] = face_value_num
        data_dict['formatted_face_value'] = formatted_face_value
        if debug:
            print(f"💵 Extracted face value: {face_value_num} ({formatted_face_value})")
    #amount raised
    if 'issue_size_number' in data_dict and 'face_value_num' in data_dict:
        amount_raised = calculate_amount_raised(data_dict)
        data_dict['amount_raised'] = amount_raised
        if debug:
            print(f"💰 Total Value: {amount_raised}")
    # Extract discount value number
    if 'Discount at which security is issued' in data_dict:
        discount_value_num, formatted_discount_value = extract_discount_value_number(data_dict['Discount at which security is issued'])
        data_dict['discount_value_num'] = discount_value_num
        data_dict['formatted_discount_value'] = formatted_discount_value
        if debug:
            print(f"💹 Extracted discount value: {discount_value_num}% ({formatted_discount_value})")
    # Extract issue price number
    if 'Issue Price' in data_dict:
        issue_price_num, formatted_issue_price = extract_issue_price_number(data_dict['Issue Price'])
        data_dict['issue_price_num'] = issue_price_num
        data_dict['formatted_issue_price'] = formatted_issue_price
        if debug:
            print(f"💲 Extracted issue price: {issue_price_num} ({formatted_issue_price})")
    # Create DataFrames
    if data_dict:
        df = pd.DataFrame([data_dict])
        if debug:
            print(f"\n📊 Created DataFrame with {len(data_dict)} columns")
    else:
        df = pd.DataFrame()
        if debug:
            print("⚠️ No data extracted - empty DataFrame created")
    
    # Write to Excel with enhanced formatting
    with pd.ExcelWriter(xlsx_file_path, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name="Terms", index=False)
        
        # Add debug sheet if requested
        if debug and debug_info:
            debug_df = pd.DataFrame(debug_info)
            debug_df.to_excel(writer, sheet_name="Debug_Log", index=False)
        
        # Enhanced formatting
        workbook = writer.book
        
        if not df.empty:
            worksheet = writer.sheets['Terms']
            
            # Header formatting
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'top',
                'bg_color': '#D7E4BC',
                'border': 1
            })
            
            # Data formatting
            data_format = workbook.add_format({
                'text_wrap': True,
                'valign': 'top',
                'border': 1
            })
            
            # Apply formatting and adjust column widths
            for col_num, column_name in enumerate(df.columns.values):
                worksheet.write(0, col_num, column_name, header_format)
                
                # Calculate appropriate column width
                max_length = max(
                    len(str(column_name)),
                    df[column_name].astype(str).map(len).max() if not df.empty else 0
                )
                adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
                worksheet.set_column(col_num, col_num, adjusted_width, data_format)
        
        # Format debug sheet if it exists
        if debug and debug_info:
            debug_worksheet = writer.sheets['Debug_Log']
            debug_header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#FFE6CC',
                'border': 1
            })
            
            for col_num, column_name in enumerate(debug_df.columns.values):
                debug_worksheet.write(0, col_num, column_name, debug_header_format)
                max_length = max(
                    len(str(column_name)),
                    debug_df[column_name].astype(str).map(len).max() if not debug_df.empty else 0
                )
                debug_worksheet.set_column(col_num, col_num, min(max_length + 2, 60))

    print(f"\n✅ Export complete!")
    print(f"📊 Total fields extracted: {len(data_dict)}")
    print(f"📁 Output file: {xlsx_file_path}")
    
    if debug:
        print(f"🔍 Debug log with {len(debug_info)} entries saved to 'Debug_Log' sheet")
        print("\n📋 Sample extracted data:")
        for i, (key, value) in enumerate(list(data_dict.items())[:10]):
            print(f"  {i+1}. {key}: {value[:100]}{'...' if len(str(value)) > 100 else ''}")
    
    return data_dict

# Enhanced usage function
# def process_document_with_validation(docx_path, xlsx_path, debug=True):
#     """
#     Process document with additional validation and error handling
#     """
#     try:
#         print(f"🔄 Starting conversion: {docx_path} → {xlsx_path}")
#         result = convert_docx_to_xlsx(docx_path, xlsx_path, debug=debug)
        
#         if result:
#             print(f"✅ Successfully extracted {len(result)} fields")
#             return result
#         else:
#             print("⚠️ No data was extracted from the document")
#             return {}
            
    # except Exception as e:
    #     print(f"❌ Error processing document: {str(e)}")
    #     return {}

# Usage examples:
"""
# Basic usage
data = convert_docx_to_xlsx('Term Sheet - Series 129 for claude.docx', 'output.xlsx')

# With comprehensive debugging
data = convert_docx_to_xlsx('Term Sheet - Series 129 for claude.docx', 'output_debug.xlsx', debug=True)

# Using the validation wrapper
data = process_document_with_validation('Term Sheet - Series 129 for claude.docx', 'validated_output.xlsx')

# Print extracted data summary
if data:
    print("\\nExtracted Fields:")
    for key, value in data.items():
        print(f"{key}: {value}")
"""

def handle_processing_date(date_string):
    """
    Process the date received from frontend and format it for Excel output
    Args:
        date_string (str): Date in YYYY-MM-DD format from frontend
    Returns:
        dict: Dictionary containing formatted date components
    """
    from datetime import datetime
    
    try:
        # Parse the date string (expected format: YYYY-MM-DD)
        parsed_date = datetime.strptime(date_string, "%Y-%m-%d")
        
        # Format into DDMMYYYY
        formatted_date = parsed_date.strftime("%d%m%Y")

        formatted_string = parsed_date.strftime("%d %B, %Y")
        
        # Create dictionary with date components
        date_components = {
            'processing_date': date_string,
            'processing_date_formatted': formatted_date,
            'processing_date_string': formatted_string,
            'D1': formatted_date[0],
            'D2': formatted_date[1],
            'M1': formatted_date[2],
            'M2': formatted_date[3],
            'Y1': formatted_date[4],
            'Y2': formatted_date[5],
            'Y3': formatted_date[6],
            'Y4': formatted_date[7]
        }
        
        return date_components
    except ValueError as e:
        print(f"⚠️ Error processing date: {e}")
        return None
    
    # Add this function near the top of the file
def extract_series_number(product_code):
    """Extract series number from product code"""
    if not product_code:
        return ""
    
    # Split by spaces and get the last part
    parts = product_code.strip().split()
    if parts:
        last_part = parts[-1]
        # If it contains 'series' (case insensitive), return the whole last part
        if 'series' in last_part.lower():
            return last_part
        # If the last part is just a number, add 'Series' prefix
        if last_part.isdigit():
            return f"Series {last_part}"
    return ""

def extract_issue_size_number(issue_size):
    """
    Extract both the first number (issue size) and fourth number (total amount) from issue size text
    Returns tuple: (issue_size_num, formatted_total_amount)
    Example input: "75 Debentures bearing face value of Rs. 1,00,000/- each, issued at Rs.95,500/-, 
                   total amounting to Rs.71,62,500/-"
    Example output: ("75", "71,62,500")
    """
    import re
    
    if not issue_size:
        return ("", "")
    
    # Remove all commas first
    cleaned_text = issue_size.replace(',', '')
    
    # Find all numbers (including decimals) in the text
    numbers = re.findall(r'\d+\.?\d*', cleaned_text)
    
    # Get first number (issue size)
    issue_size_num = numbers[0] if numbers else ""
    
    # Get fourth number if available and format it
    formatted_amount = ""
    if len(numbers) >= 4:
        total_amount = numbers[3]
        try:
            amount = int(float(total_amount))
            formatted_amount = format_indian_number(amount)
        except (ValueError, TypeError):
            pass
    
    return (issue_size_num, formatted_amount)


def extract_tenor_days_number(tenor_text):
    """Extract number from Tenor in Days text"""
    import re
    
    if not tenor_text:
        return ""
    
    # Find numbers in the text
    numbers = re.findall(r'\d+', tenor_text)
    
    # Return first number found or empty string if none found
    return numbers[0] if numbers else ""

def extract_face_value_number(face_value):
    """
    Extract number from Face Value text and return both raw and formatted values
    Returns tuple: (number_as_string, formatted_with_commas)
    Example: "Rs. 1,00,000/- (Rupees One Lakh Only)" -> ("100000", "1,00,000")
    """
    import re
    
    if not face_value:
        return ("", "")
    
    # Remove commas from numbers
    cleaned_text = face_value.replace(',', '')
    
    # Find numbers including decimals
    numbers = re.findall(r'\d+\.?\d*', cleaned_text)
    
    if numbers:
        try:
            # Convert to integer if possible
            number_str = numbers[0]
            number_int = int(float(number_str))
            formatted = format_indian_number(number_int)
            return (str(number_int), formatted)
        except (ValueError, TypeError):
            return (numbers[0], numbers[0])
    
    return ("", "")

def calculate_amount_raised(data_dict):
    """
    Calculate total value by multiplying issue size and face value,
    converting from lakhs to crores
    """
    try:
        # Get values and convert to float
        issue_size = float(data_dict.get('issue_size_number', 0))
        face_value = float(data_dict.get('face_value_num', 0))
        
        # Calculate total in lakhs
        total_in_lakhs = issue_size * face_value
        
        # Convert to crores (1 crore = 100 lakhs)
        total_in_crores = total_in_lakhs / 10000000
        
        # Format with comma separators and 2 decimal places
        formatted_value = f"Rs {total_in_crores:,.2f} crores"
        
        return formatted_value
    except (ValueError, TypeError) as e:
        print(f"⚠️ Error calculating total value: {e}")
        return "Rs 0 crores"

def extract_discount_value_number(discount_text):
    """
    Extract first complete number from discount text, handling commas.
    Returns a tuple: (number_as_string, formatted_with_commas)
    Example: "Rs. 4,500/- (Rupees Four Thousand Five Hundred Only)" -> ("4500", "4,500")
    """
    import re

    if not discount_text:
        return ("", "")

    # First find any number pattern that may include commas
    number_pattern = r'(?:Rs\.?\s*)?(\d+(?:,\d+)*(?:\.\d+)?)'
    matches = re.search(number_pattern, discount_text)

    if matches:
        # Remove commas from the matched number
        number_str = matches.group(1).replace(',', '')
        try:
            # Try to convert to int if possible, else float
            if '.' in number_str:
                number_val = float(number_str)
            else:
                number_val = int(number_str)
            formatted = format_indian_number(int(float(number_str)))
            return (number_str, formatted)
        except (ValueError, TypeError):
            return (number_str, number_str)

    return ("", "")

def extract_issue_price_number(issue_price_text):
    """
    Extract first complete number from Issue Price text, handling commas.
    Returns a tuple: (number_as_string, formatted_with_commas)
    Example: "Rs. 95,500/- (Rupees Ninety-Five Thousand Five Hundred Only) per Debenture"
    Output: ("95500", "95,500")
    """
    import re

    if not issue_price_text:
        return ("", "")

    # Find first number pattern that may include commas
    number_pattern = r'(?:Rs\.?\s*)?(\d+(?:,\d+)*(?:\.\d+)?)'
    match = re.search(number_pattern, issue_price_text)
    if match:
        number_str = match.group(1).replace(',', '')
        try:
            number_int = int(float(number_str))
            formatted = format_indian_number(number_int)
            return (number_str, formatted)
        except (ValueError, TypeError):
            return (number_str, number_str)
    return ("", "")

def format_indian_number(number):
    """Helper function to format number with Indian style commas"""
    str_num = str(number)
    length = len(str_num)
    
    if length <= 3:
        return str_num
    
    # Split the number into parts
    last_three = str_num[-3:]
    remaining = str_num[:-3]
    
    # Add commas every 2 digits in the remaining part
    formatted = ''
    for i, digit in enumerate(reversed(remaining)):
        if i % 2 == 1 and i != len(remaining) - 1:
            formatted = ',' + digit + formatted
        else:
            formatted = digit + formatted
    
    return formatted + ',' + last_three

def save_uploaded_file(file, upload_folder):
    import os
    file_path = os.path.join(upload_folder, file.filename)
    file.save(file_path)
    return file_path